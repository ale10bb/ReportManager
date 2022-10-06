# -*- coding: UTF-8 -*-
import os
import logging
import sys
import argparse
import re
from email.utils import parseaddr
from walkdir import filtered_walk, file_paths
from docx import Document
import requests
import base64
from . import archive, mail, mysql
from . import var

# --------------------------------------------
#           检查邮件及附件的处理逻辑
# --------------------------------------------
# 注：仅可由命令逻辑调用，默认参数均为合法

def check_mail_content(mail_content:dict) -> dict:
    ''' 读取{mail_content}中的内容，获取发件人和指令，处理完毕时返回警告信息及处理结果'content'。

    Args:
        mail_content(dict): check_result中解析后的邮件内容；

    Returns:
        {'warnings': [], 'content': {'user_id': (str), 'name': (str), 'urgent': (bool), 'excludes': (list), 'force': (str)}}

    Raises:
        ValueError: 如果发件人非法
    '''
    logger = logging.getLogger(__name__)
    logger.debug('args: {}'.format({'mail_content': mail_content}))
    assert type(mail_content) == dict, 'invalid arg: mail_content'

    ret = {'warnings': [], 'content': {'user_id': '', 'name': '', 'urgent': False, 'excludes': [], 'force': ''}}
    # 从from中读取发信人并校验，校验内容为：
    #   域名必须为配置文件中的_MAIL_DOMAIN & 用户名必须位于user表中
    #   如果subject中包含"--sender userid"参数，则在检验userid有效后，将其作为发信人
    # 预期结果：
    #   流程正常完成时，在ret中填入user_id、name
    #   校验失败时，抛出ValueError，包含发件人邮箱
    parser = argparse.ArgumentParser()
    parser.add_argument(
        '-s',
        '--sender',
        nargs='?',
        const='',
        default='',
        action='store'
    )
    parser.add_argument('others', nargs='*')
    args = parser.parse_args(re.sub(' +', ' ',mail_content['subject'].strip()).split(' '))
    if args.sender:
        user_id = args.sender
        logger.info('manual sender: {}'.format(args.sender))
    else:
        user_id = parseaddr(mail_content['from'])[1].split('@')[0]

    if mail.valid_domain(mail_content['from']) and mysql.t_user.__contains__(user_id):
        name = mysql.t_user.fetch(user_id)[1]
        ret['content']['user_id'] = user_id
        ret['content']['name'] = name
        logger.info('sender: {}/{}'.format(user_id, name))
    else:
        raise ValueError('Invalid sender "{}"'.format(mail_content['from']))

    # 从content中指令并校验，校验内容为：
    #   是否输入“加急” & 输入内容是否合法
    #   是否输入“组员” & 组员是否具备审核资格
    #   是否输入“指定” & 是否同时设置了“加急” & 指定是否具备审核资格 & 指定是否为发件人或组员
    # 预期结果：
    #   流程正常完成时，在ret中填入urgent、excludes、force。非法指令将被忽略并使用默认值（见下）
    urgent = False
    excludes = []
    force = ''
    content = ''.join(mail_content['content'])
    # logger输出时，\r\n将被渲染导致日志换行，因此替换成<p>增加可读性
    logger.debug('content: {}'.format(re.sub('[\r\n]+', '<p>', content)))
    # 循环读取content中的所有行，尝试校验指令
    # 注意：如果指令重复，将丢弃上一条结果，将结果重置为默认值并处理新行
    for line in content.splitlines():
        # 有人习惯打冒号、全角逗号、半角逗号，一并替换成合法值
        line = re.sub(' +|：', ' ', line.strip())
        line = re.sub('，|,', '、', line.strip())
        # 替换常见错误输入后，按第一个空格分隔1次，长度不为2的指令行将被忽略
        cmds = line.split(' ', 1)
        if len(cmds) != 2:
            continue
        logger.debug('cmds: {}'.format(cmds))
        if cmds[0] == '加急':
            urgent = False
            if cmds[1] == '是' or cmds[1] == '1':
                urgent = True
                logger.info('urgent: {}'.format(urgent))
            elif cmds[1] == '否' or cmds[1] == '0':
                urgent = False
            else:
                ret['warnings'].append('无效的"加急"值，已忽略')
            continue
        if cmds[0] == '组员':
            excludes = []
            members = cmds[1].split('、')
            # 检查组员并修改为user_id
            for member in cmds[1].split('、'):
                excludes.extend(
                    mysql.t_user.search(user_id=member, only_reviewer=True)['user'] +
                    mysql.t_user.search(name=member, only_reviewer=True)['user']
                )
            excludes = [i[0] for i in excludes]
            if len(members) != len(excludes):
                ret['warnings'].append('已去除无效组员 {} -> {}'.format(members, excludes))
            logger.info('excludes: {}'.format(excludes))
        if cmds[0] == '指定':
            # 检查指定并修改为主键
            users = mysql.t_user.search(user_id=cmds[1], only_reviewer=True)['user'] + mysql.t_user.search(name=cmds[1], only_reviewer=True)['user']
            if len(users) == 1:
                force = users[0][0]
            else:
                ret['warnings'].append('已去除无效指定 "{}"'.format(cmds[1]))
            logger.info('force: {}'.format(force))

    # 考虑到指令重复、指令之间有因果关系等情况，仅当读取完所有行之后，再进行总体校验及写入ret
    # “指定”回滚逻辑：
    #   指定了自己或组员
    if force == user_id or force in excludes:
        force = ''
        ret['warnings'].append('指定失败: 项目相关人员')
        logger.warning('rallbacked force due to "in excludes"')

    # 写入ret
    ret['content']['urgent'] = urgent
    ret['content']['excludes'] = excludes
    ret['content']['force'] = force
    logger.debug('return: {}'.format(ret))
    return ret


def check_mail_attachment(work_path:str, operator:str) -> dict:
    ''' 读取存放在{work_path}中的附件，根据操作符获取文档信息。处理完毕后，返回attachment信息。

    Args:
        work_path(str)：工作目录（存放附件的目录）
        operator(str): 操作符（submit/finish）

    Returns:
        {'warnings': [], 'attachment': {'pages': (int), 'company': (str), 'codes': (list), 'names': (dict)}}

    Raises:
        AssertionError: 如果参数类型非法
        ValueError: 如果未读取到有效文档
    '''
    logger = logging.getLogger(__name__)
    logger.debug('args: {}'.format({'work_path': work_path, 'operator': operator}))
    assert operator == 'submit' or operator == 'finish', 'invalid arg: operator'
    assert os.path.isdir(work_path), 'invalid arg: work_path'

    ret = {'warnings': [], 'attachment': {}}
    # 解压工作目录中的所有压缩包，并添加解压失败的告警
    for archive_file_path_with_error in archive.extract(work_path):
        ret['warnings'].append('解压失败："{}"'.format(os.path.basename(archive_file_path_with_error)))

    # 读取工作目录中的所有文档
    # 传入的操作符用于控制读取逻辑
    # 用返回值的codes参数作为标记，无codes说明读取失败，此时抛出ValueError异常
    if operator == 'submit':
        ret['attachment'] = read_document(work_path)
        if not ret['attachment']['codes']:
            raise ValueError('No valid documents within "submit"')
    if operator == 'finish':
        ret['attachment'] = read_XT13(work_path)
        if not ret['attachment']['codes']:
            raise ValueError('No valid documents within "finish"')

    logger.debug('return: {}'.format(ret))
    return ret

# -------------------------------------------
#           检查文档内容的处理逻辑
# -------------------------------------------
# 注：仅可由命令逻辑调用，默认参数均为合法
#     函数传入work_path，对工作目录下的所有文件进行处理

def read_document(work_path:str) -> dict:
    ''' 读取{work_path}下的所有word文档，读取项目编号、项目名称、委托单位、文档页数，并尝试将文档转换到最新格式。

    Args:
        work_path: 工作目录

    Returns:
        dict: {'pages': 0, 'company': '', 'codes': [], 'names': {}}
    '''
    logger = logging.getLogger(__name__)
    logger.debug('args: {}'.format({'work_path': work_path}))
    assert os.path.isdir(work_path), 'invalid arg: work_path'

    # 一个work_path作为一个项目包，读取时会将所有有效文档的页数相加，并返回项目名称的集合
    # 但默认项目包里面只有一个委托单位
    ret = {'pages': 0, 'company': '', 'codes': [], 'names': {}}
    # work_path中的所有doc和docx文件视作有效文档，但忽略审核意见单和临时文件
    document_paths = file_paths(filtered_walk(work_path, included_files=['*.doc', '*.docx'], excluded_files=['~$*', '*XT13*', '*签发意见单*']))
    for document_path in document_paths:
        logger.info('reading "{}"'.format(os.path.basename(document_path)))
        # 初始化读取变量
        page = 30
        company = ''
        code = ''
        name = ''
        
        # 尝试转换文件及读取页数
        try:
            if sys.platform == 'win32':
                import win32com.client, pythoncom
                pythoncom.CoInitialize()
                word = win32com.client.gencache.EnsureDispatch('Word.Application')
                document = word.Documents.Open(FileName=document_path)
                logger.debug('compatibility mode: {}'.format(document.CompatibilityMode))
                # 将word转换为最新文件格式
                if document.CompatibilityMode < 15: # CompatibilityMode=11-14(旧版)
                    document.Convert()
                    document.Save()
                    document_path = os.path.splitext(document_path)[0] + '.docx'
                    logger.info('converted to latest compatibility mode')
                # 读取文档页数
                page = document.ComputeStatistics(2) # wdStatisticPages=2
                logger.info('page: {}'.format(page))
                document.Close(SaveChanges=0)
                pythoncom.CoUninitialize()
            else:
                assert var.dedicate_win32, 'win32_handler not set'
                with open(document_path,'rb') as f:
                    files = {'document': f}
                    r = requests.post(var.dedicate_win32, files=files, timeout=60).json()
                assert not r['result'], r['err']
                page = r['data']['page']
                logger.info('page: {}'.format(page))
                if r['data']['converted']['name']:
                    logger.info('converted to latest compatibility mode')
                    os.remove(document_path)
                    with open(os.path.join(os.path.dirname(document_path), r['data']['converted']['name']), 'wb') as f:
                        f.write(base64.b64decode(r['data']['converted']['content'].encode('utf-8')))
                    document_path = os.path.join(os.path.dirname(document_path), r['data']['converted']['name'])
        except:
            logger.error('win32_handler error', exc_info=True)

        # 调python-docx读取信息
        # 如果读取失败，直接跳过
        try:
            document = Document(document_path)
        except:
            logger.error('ignored document', exc_info=True)
            continue
        
        # 读项目编号
        ## 印象中所有项目编号都能在前几行读到
        for paragraph in document.paragraphs[0:5]:
            re_result = re.search('SHTEC20[0-9]{2}(PRO|PST|DSYS|SOF|SRV|PER|FUN)[0-9]{4}([-_][0-9]+){0,1}', paragraph.text)
            if re_result:
                code = re_result[0]
                break
        logger.info('code: {}'.format(code))
        ## 读不到项目编号时，忽略该文档
        if not code:
            logger.warning('ignored document')
            continue

        # 特别地，附件和复核意见单不把编号和名称写入ret，但计算页数
        flag = False
        ## 注：以下所有[0:30]表示读取封面页，但最多读取文档的前30行
        ## 考虑到“正常”页边距的一页文档可包含40行5号字段落，且封面页一定不会塞满40行5号字，因此读到30行就认为超过了封面
        for paragraph in document.paragraphs[0:30]:
            # 检测封面是否包含这两种文档的关键词
            if re.search('附 {0,10}件', paragraph.text) or re.search('复核意见表', paragraph.text):
                flag = True
                break
        if flag:
            logger.info('ignored but counted pages')
            ret['pages'] = ret['pages'] + page
            continue

        # 读取系统名称和委托单位
        ## 等保报告
        if 'DSYS' in code:
            logger.debug('reading DSYS')
            # 系统名称在封面页的表格外面
            for paragraph in document.paragraphs[0:30]:
                if '等级测评报告' in paragraph.text:
                    name = paragraph.text.split('等级测评报告', 1)[0].strip()
                    break
            # 从表格中读取委托单位
            for row in document.tables[0].rows:
                if '单位' in row.cells[0].text and not '测评单位' in row.cells[0].text:
                    company = row.cells[1].text.strip()
                    break
        ## PRO和PST
        elif 'PRO' in code or 'PST' in code or 'PER' in code:
            logger.debug('reading PRO/PST/PER')
            # 直接从第一个表格中读取
            for row in document.tables[0].rows:
                if '名称' in row.cells[0].text:
                    name = row.cells[1].text.strip()
                if '委托单位' in row.cells[0].text:
                    company = row.cells[1].text.strip()
        ## SOF
        elif 'SOF' in code or 'FUN' in code:
            logger.debug('reading SOF/FUN')
            # 遍历第一页的行读取
            for paragraph in document.paragraphs[0:30]:
                if '名称' in paragraph.text:
                    name = re.sub('^.*名称(:|：)', '', paragraph.text).strip()
                if '委托单位' in paragraph.text:
                    company = re.sub('^.*单位(:|：)', '', paragraph.text).strip()
                if name and company:
                    break
        ## 其他报告，自求多福
        else:
            logger.debug('reading others')
            name = '<unknown>'
            # 先尝试读表格
            for row in document.tables[0].rows:
                if '名称' in row.cells[0].text:
                    name = row.cells[1].text.strip()
                if '委托单位' in row.cells[0].text:
                    company = row.cells[1].text.strip()
            # 再尝试读行
            for paragraph in document.paragraphs[0:30]:
                if '名称' in paragraph.text:
                    name = re.sub('^.*名称(:|：)', '', paragraph.text).strip()
                if '委托单位' in paragraph.text:
                    company = re.sub('^.*单位(:|：)', '', paragraph.text).strip()
                if name and company:
                    break

        # 尝试去除可能存在的换行符
        name = name.replace('\n', '')
        company = company.replace('\n', '')
        logger.info('name: {}'.format(name))
        logger.info('company: {}'.format(company))
        # 为了防止“跳过附件和复核意见单失败”的情况，仅读取到有效值后才把编号和名称写入ret
        if name:
            ret['codes'].append(code)
            ret['names'][code] = name
        # 委托单位有效时覆盖缓存值
        if company:
            ret['company'] = company
        # 累加项目包总页数
        ret['pages'] = ret['pages'] + page

    logger.debug('return: {}'.format(ret))
    return ret


def read_XT13(work_path:str) -> dict:
    ''' 读取{work_path}下的所有审核意见单，读取项目编号。

    Args:
        work_path: 工作目录

    Returns:
        dict: {'pages': 0, 'company': '', 'codes': (list), 'names': {}}
    '''
    logger = logging.getLogger(__name__)
    logger.debug('args: {}'.format({'work_path': work_path}))
    assert os.path.isdir(work_path), 'invalid arg: work_path'

    ret = {'pages': 0, 'company': '', 'codes': [], 'names': {}}
    # 只读docx格式的审核意见单
    document_paths = file_paths(filtered_walk(work_path, included_files=['*XT13*.docx', '*签发意见单*.docx'], excluded_files=['~$*']))
    for document_path in document_paths:
        logger.info('reading "{}"'.format(os.path.basename(document_path)))
        document = Document(document_path)
        # 从审核意见单第一行读编号
        re_result = re.search('SHTEC20[0-9]{2}(PRO|PST|DSYS|SOF|SRV|PER|FUN)[0-9]{4}([-_][0-9]+){0,1}', document.paragraphs[0].text)
        if re_result:
            ret['codes'].append(re_result[0])
        logger.info('code: {}'.format(re_result[0]))
                
    logger.debug('return: {}'.format(ret))
    return ret
