# -*- coding: UTF-8 -*-
""" 文档操作工具类
"""
import os
import shutil
import logging
import re
from walkdir import filtered_walk, file_paths
import win32com.client
from docx import Document

from .types import *


def read_document(work_path: str) -> Attachment:
    """读取{work_path}下的所有word文档，读取项目编号、项目名称、委托单位、文档页数

    Args:
        work_path(str): 工作目录

    Returns:
        Attachment
    """
    logger = logging.getLogger(__name__)
    logger.debug("args: %s", {"work_path": work_path})

    # 一个work_path作为一个项目包，读取时会将所有有效文档的页数相加，并返回项目名称的集合
    # 但默认项目包里面只有一个委托单位
    ret: Attachment = {"pages": 0, "company": "", "names": {}}
    # work_path中的所有doc和docx文件视作有效文档，但忽略审核意见单和临时文件
    document_paths = file_paths(
        filtered_walk(
            work_path,
            included_files=["*.doc", "*.docx"],
            excluded_files=["~$*", "*XT13*", "*签发意见单*"],
        )
    )
    word = win32com.client.gencache.EnsureDispatch("Word.Application")
    for document_path in document_paths:
        logger.info('reading "%s"', os.path.basename(document_path))
        page = 0
        code = ""
        name = ""
        company = ""
        document = None
        try:
            document = word.Documents.Open(FileName=document_path)
            # 读取文档页数
            page = document.ComputeStatistics(2)  # wdStatisticPages=2
            logger.info("page: %s", page)
            # 读项目编号
            # 印象中所有项目编号都能在前几行读到
            pattern = re.compile(
                "SHTEC20[0-9]{2}(PRO|PST|DSYS|SOF|SRV|PER|FUN|PCT)[0-9]{4}([-_][0-9]+){0,1}"
            )
            for i in range(5):
                re_result = re.search(pattern, document.Paragraphs(i + 1).Range.Text)
                if re_result:
                    code = re_result.group()
                    logger.info("code: %s", code)
                    break
            else:
                logger.warning("ignored document")
                continue
            # 附件和复核意见单的逻辑已去除

            # 读取系统名称和委托单位
            if "DSYS" in code:
                logger.debug("reading DSYS")
                # 从基本信息表中读取
                name = document.Tables(2).Cell(2, 2).Range.Text
                company = document.Tables(2).Cell(5, 2).Range.Text
            elif "PRO" in code or "PST" in code or "PER" in code or "PCT" in code:
                logger.debug("reading PRO/PST/PER/PCT")
                # 直接从第一个表格中读取
                name = document.Tables(1).Cell(1, 2).Range.Text
                company = document.Tables(1).Cell(2, 2).Range.Text
            elif "SOF" in code or "FUN" in code:
                logger.debug("reading SOF/FUN")
                # 遍历第一页的行读取
                for i in range(30):
                    paragraph = document.Paragraphs(i + 1).Range.Text.strip()
                    if "名称" in paragraph:
                        name = re.sub("^.*名称(:|：)", "", paragraph)
                    if "委托单位" in paragraph:
                        company = re.sub("^.*单位(:|：)", "", paragraph)
                    if name and company:
                        break
            # 其他报告，自求多福
            else:
                logger.debug("reading others")
                # 先尝试读表格
                for i in range(document.Tables(1).Rows.Count):
                    if "项目名称" in document.Tables(1).Cell(i + 1, 1).Range.Text:
                        name = document.Tables(1).Cell(i + 1, 2).Range.Text
                    elif "报告名称" in document.Tables(1).Cell(i + 1, 1).Range.Text:
                        name = document.Tables(1).Cell(i + 1, 2).Range.Text
                    elif "系统名称" in document.Tables(1).Cell(i + 1, 1).Range.Text:
                        name = document.Tables(1).Cell(i + 1, 2).Range.Text
                    elif "委托单位" in document.Tables(1).Cell(i + 1, 1).Range.Text:
                        company = document.Tables(1).Cell(i + 1, 2).Range.Text
                    elif "被测单位" in document.Tables(1).Cell(i + 1, 1).Range.Text:
                        company = document.Tables(1).Cell(i + 1, 2).Range.Text
                # 再尝试读行
                for i in range(30):
                    paragraph = document.Paragraphs(i + 1).Range.Text.strip()
                    if (
                        "项目名称" in paragraph
                        or "报告名称" in paragraph
                        or "系统名称" in paragraph
                    ):
                        name = re.sub("^.*名称(:|：)", "", paragraph)
                    elif "委托单位" in paragraph or "被测单位" in paragraph:
                        company = re.sub("^.*单位(:|：)", "", paragraph)
                    if name and company:
                        break
            # 尝试去除可能存在的换行符
            if name:
                name = re.sub("(\r|\n|\x07| *)", "", name)
                logger.info("name: %s", name)
                ret["names"][code] = name
            # 委托单位有效时覆盖缓存值
            if company:
                company = re.sub("(\r|\n|\x07| *)", "", company)
                logger.info("company: %s", company)
                ret["company"] = company
            # 累加项目包总页数
            ret["pages"] = ret["pages"] + page
        except Exception:
            logger.warning("read failed", exc_info=True)
        finally:
            if document:
                document.Close(SaveChanges=0)

    logger.debug("return: %s", ret)
    return ret


def read_XT13(work_path: str) -> Attachment:
    """读取{work_path}下的所有审核意见单，读取项目编号

    Args:
        work_path: 工作目录

    Returns:
        Attachment
    """
    logger = logging.getLogger(__name__)
    logger.debug("args: %s", {"work_path": work_path})

    ret: Attachment = {"pages": 0, "company": "", "names": {}}
    # 只读docx格式的审核意见单
    document_paths = file_paths(
        filtered_walk(
            work_path,
            included_files=["*XT13*.docx", "*MP07*.docx", "*签发意见单*.docx"],
            excluded_files=["~$*"],
        )
    )
    word = win32com.client.gencache.EnsureDispatch("Word.Application")
    for document_path in document_paths:
        logger.info('reading "%s"', os.path.basename(document_path))
        code = ""
        name = ""
        document = None
        try:
            document = word.Documents.Open(FileName=document_path)
            re_result = re.search(
                "SHTEC20[0-9]{2}(PRO|PST|DSYS|SOF|SRV|PER|FUN|PCT)[0-9]{4}([-_][0-9]+){0,1}",
                document.Paragraphs(1).Range.Text,
            )
            if re_result:
                code = re_result.group()
                logger.info("code: %s", code)
                name = document.Tables(1).Cell(2, 2).Range.Text
                name = re.sub("(\r|\n|\x07| *)", "", name)
                logger.info("name: %s", name)
                ret["names"][code] = name
        except Exception:
            logger.warning("read failed", exc_info=True)
        finally:
            if document:
                document.Close(SaveChanges=0)

    word.Quit()
    logger.debug("return: %s", ret)
    return ret


def gen_XT13(authorname: str, code: str, project_name: str, target_path):
    """生成XT13

    Args:
        authorname(str): 撰写人名字
        code(str): 项目编号
        project_name(str): 项目名称
        target_path(str): 保存路径
    """
    logger = logging.getLogger(__name__)
    logger.debug(
        "args: %s",
        {
            "authorname": authorname,
            "code": code,
            "project_name": project_name,
            "target_path": target_path,
        },
    )

    logger.info('generating XT13 for "%s" to "%s"', code, target_path)
    try:
        # 根据项目编号读取不同的审核意见单模板
        if "PRO" in code or "PST" in code:
            template_document = Document(
                os.path.join("template", "RD-XT13测评报告审核、签发意见单-PROPST.docx")
            )
        elif "SOF" in code:
            template_document = Document(
                os.path.join("template", "RD-XT13测评报告审核、签发意见单-SOF.docx")
            )
        elif "DSYS" in code:
            template_document = Document(
                os.path.join("template", "RD-XT13测评报告审核、签发意见单-DSYS.docx")
            )
        elif "SRV" in code:
            template_document = Document(
                os.path.join("template", "RD-XT13测评报告审核、签发意见单-SRV.docx")
            )
        elif "PER" in code:
            template_document = Document(
                os.path.join("template", "RD-XT13测评报告审核、签发意见单-PER.docx")
            )
        elif "FUN" in code:
            template_document = Document(
                os.path.join("template", "RD-XT13测评报告审核、签发意见单-FUN.docx")
            )
        else:
            # 目前无视其他类型的报告
            return
        # 项目编号(bold)
        template_document.paragraphs[0].add_run(code).bold = True
        # 项目名称
        template_document.tables[0].cell(1, 1).text = project_name
        # 报告撰写人
        template_document.tables[0].cell(2, 1).text = authorname
        template_document.save(target_path)
    except:
        logger.warning("generating XT13 failed", exc_info=True)


def encrypt(document_path: str):
    """加密文档（在安装DLP的worker上直接另存为）

    Args:
        document_path(str): 文档路径
    """
    logger = logging.getLogger(__name__)
    logger.debug("args: %s", {"document_path": document_path})

    word = win32com.client.gencache.EnsureDispatch("Word.Application")
    document = None
    try:
        logger.info('encrypting "%s"', os.path.basename(document_path))
        document = word.Documents.Open(FileName=document_path)
        document.SaveAs2(FileName=document_path + ".enc")
        document.Close(SaveChanges=0)
        document = None
        os.remove(document_path)
        shutil.move(document_path + ".enc", document_path)
    except:
        logger.warning("Encryption failed.", exc_info=True)
    finally:
        if document:
            document.Close(SaveChanges=0)
