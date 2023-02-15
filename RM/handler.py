# -*- coding: UTF-8 -*-
import os
import logging
import shutil
import datetime
import json
import multiprocessing
from walkdir import filtered_walk, file_paths, dir_paths
from docx import Document
from . import archive, dingtalk, mail, mysql, wxwork
from . import notification, validator
from . import var

def do_attend():
    ''' 打卡提醒函数的主入口。调用后向主通知群发送打卡提示，包含当前任务、分配队列和交互入口。向企业微信发送任务提醒。
    '''
    logger = logging.getLogger(__name__)
    
    # 准备通知所需数据
    currents = mysql.t_current.search(page_size=9999)['all']
    # 24小时没有动作的情况下跳过信息输出
    if not currents and (datetime.datetime.now().timestamp() - mysql.t_history.pop()[6] > 86400):
        logger.debug('skipped output')
        return
    logger.info('currents: {}'.format(currents))
    queue = mysql.t_user.pop(count=9999, hide_busy=False)
    logger.info('queue: {}'.format(queue))
    currents_group_by_user_id = {}
    for user_id in [item[0] for item in queue]:
        currents_group_by_user_id[user_id] = []
    for project in currents:
        currents_group_by_user_id[project[3]].append('+'.join(json.loads(project[10]).keys()))

    # 钉钉当前项目
    lines = []
    for project in currents:
        delta = datetime.datetime.now() - datetime.datetime.fromtimestamp(project[5])
        lines.append('- {}{} -> {} ({})'.format(
            '+'.join(json.loads(project[10]).keys()), 
            '/急' if project[8] else '', 
            project[4], 
            '{}d{}h'.format(delta.days, int(delta.seconds / 3600))
        ))
    part1 = '**当前审核任务**\n\n' + '\n'.join(lines) if lines else '**当前无审核任务**'
    # 钉钉分配队列
    part2 = '**分配队列**\n\n' + '\n'.join([
        '1. {}{}'.format(
            row[1],
            ' (+{})'.format(row[5]) if row[5] else ''
        ) for row in queue
    ])
    dingtalk.send_action_card(part1+'\n\n---\n\n'+part2, to_stdout=var.debug)

    # 微信个人通知
    for idx, item in enumerate(queue):
        if item[3] == 0:
            status = '空闲'
        elif item[3] == 1:
            status = '不审加急'
        elif item[3] == 2:
            status = '不审报告'
        else:
            status = '未知'
        content = '===== 状态通知 =====\n\n你的状态: {}\n你的分配顺位: {}{}'.format(
            status, 
            idx + 1 if item[6] == 0 else '跳过一篇', 
            ' (+{}页)'.format(item[4]) if item[4] else ''
        )
        if currents_group_by_user_id[item[0]]:
            content += '\n你当前有{}个审核任务:\n'.format(item[5]) + '\n'.join(currents_group_by_user_id[item[0]])
        wxwork.send_text(content, to=[item[0]], to_stdout=var.debug)


def do_mail(keywords:dict=None, from_file:str=''):
    ''' 邮件管理函数的主逻辑，对于传入的单个邮件，依次进行邮件检查、数据库操作、发送通知步骤。

    传入的{keywords}用于自定义处理逻辑。传入{from_file}时，不通过POP3拉取邮件，而是读取本地eml文件。

    Args:
        keywords(dict): 拉取邮件时使用的关键词
        from_file(str): 本地eml文件的路径

    Returns:
        处理邮件数量
    '''
    logger = logging.getLogger(__name__)
    logger.debug('args: {}'.format({'keywords': keywords, 'from_file': from_file}))

    received_mails = mail.receive(os.path.join(var.storage, 'temp'), keywords, from_file)
    for check_results in received_mails:
        logger.info('reading "{}" from "{}"'.format(check_results['operator'], check_results['mail']['from']))
        try:
            # Step 1: 检查邮件内容及附件内容，得到work_path和所有相关信息，生成需要的文件并清理文件夹
            attachments_path = os.path.join(check_results['work_path'], 'attachments')

            # content的常规报错情况为：
            #   发件人非法
            ret = validator.check_mail_content(check_results['mail'])
            check_results['content'] = ret['content']
            check_results['warnings'] += ret['warnings']

            # attachment的常规报错情况为：
            #   未从附件中读取到有效文档
            ret = validator.check_mail_attachment(attachments_path, check_results['operator'])
            check_results['attachment'] = ret['attachment']
            check_results['warnings'] += ret['warnings']

            # 提交审核时，生成XT13
            if check_results['operator'] == 'submit':
                for code in check_results['attachment']['codes']:
                    try:
                        logger.info('generating XT13 for "{}"'.format(code))
                        # 根据项目编号读取不同的审核意见单模板
                        if 'PRO' in code or 'PST' in code:
                            template_document = Document(os.path.join('template', 'RD-XT13测评报告审核、签发意见单-PROPST.docx'))
                        elif 'SOF' in code:
                            template_document = Document(os.path.join('template', 'RD-XT13测评报告审核、签发意见单-SOF.docx'))
                        elif 'DSYS' in code:
                            template_document = Document(os.path.join('template', 'RD-XT13测评报告审核、签发意见单-DSYS.docx'))
                        elif 'SRV' in code:
                            template_document = Document(os.path.join('template', 'RD-XT13测评报告审核、签发意见单-SRV.docx'))
                        elif 'PER' in code:
                            template_document = Document(os.path.join('template', 'RD-XT13测评报告审核、签发意见单-PER.docx'))
                        elif 'FUN' in code:
                            template_document = Document(os.path.join('template', 'RD-XT13测评报告审核、签发意见单-FUN.docx'))
                        else:
                            # 目前无视其他类型的报告
                            continue 

                        # 项目编号(bold)
                        template_document.paragraphs[0].add_run(code).bold = True
                        # 项目名称
                        template_document.tables[0].cell(1, 1).text = check_results['attachment']['names'].get(code, '')
                        # 报告撰写人
                        template_document.tables[0].cell(2, 1).text = check_results['content']['name']

                        # 另存为
                        template_document.save(os.path.join(attachments_path, 'RD-XT13测评报告审核、签发意见单{}.docx'.format(code)))
                    except Exception as err:
                        logger.warning('generating XT13 failed: {}'.format(err))
                        check_results['warnings'].append('generating XT13 failed for "{}"'.format(code))

            # 完成审核时，删除多余文件
            if check_results['operator'] == 'finish':
                for file_path in file_paths(filtered_walk(attachments_path, excluded_files=['*.doc', '*.docx', '*.rar', '*.zip', '*.7z', '*.tar'])):
                    os.remove(file_path)
                    logger.info('removed "{}"'.format(os.path.basename(file_path)))


            # Step 2: 进行数据库操作，得到操作结果target

            # 提交审核的情况
            if check_results['operator'] == 'submit':
                # 向current库中插入该项目，add接口只需要传入项目基本信息
                current_id = mysql.t_current.add(
                    check_results['attachment']['names'], 
                    check_results['attachment']['company'],
                    check_results['attachment']['pages'],
                    check_results['content']['urgent'],
                    check_results['content']['user_id'], 
                    check_results['mail']['timestamp'] 
                )

                # 使用edit接口分配项目
                ## 如果有指定审核人，直接设置为该人（前一步骤已经校验了有效性）
                if check_results['content']['force']:
                    mysql.t_current.edit(current_id, reviewerid=check_results['content']['force'])
                ## 否则调用mysql.t_user.pop获取下一顺位的审核人，传入的excludes拼接了组员和自己
                else:
                    mysql.t_current.edit(
                        current_id, 
                        reviewerid=mysql.t_user.pop(
                            excludes=check_results['content']['excludes'] + [check_results['content']['user_id']],
                            urgent=check_results['content']['urgent'],
                        )[0][0]
                    )

                # 从current中取回操作结果，写入target
                check_results['target'] = mysql.t_current.fetch(current_id)
                logger.info('(submit) "{}" -> "{}"'.format(check_results['target'][1], check_results['target'][3]))

            # 完成审核的情况
            if check_results['operator'] == 'finish':
                codes = '+'.join(check_results['attachment']['codes'])
                ret = mysql.t_current.search('+'.join(check_results['attachment']['codes']))['all']
                assert len(ret) == 1, 'Not unique result for code: "{}".'.format(codes)
                # 从current库中正常删除项目，delete接口会自动将结果写入history
                mysql.t_current.delete(ret[0][0], check_results['mail']['timestamp'])

                # 从history中取回操作结果，写入target
                ## 由于插入history时不检测唯一性，当用户输错项目编号时，search结果可能重复。
                ## 在此取第一个结果作为有效返回
                check_results['target'] = mysql.t_history.search(codes)['all'][0]
                logger.info('(finish) "{}" <- "{}"'.format(check_results['target'][1], check_results['target'][3]))
            
            logger.debug('target: {}'.format(check_results['target']))


            # Step 3: 存档work_path，并发送通知消息
            # 该步骤主要用于生成通知内容并发送通知。邮件发送失败将抛出异常，钉钉通知发送失败将告警并继续处理

            ## 去除所有子目录结构
            for file_path in file_paths(filtered_walk(attachments_path, min_depth=1)):
                shutil.copy(file_path, attachments_path)
            for dir_path in dir_paths(filtered_walk(attachments_path, min_depth=1, depth=1)):
                shutil.rmtree(dir_path)

            ## 对于提交审核，仅将temp中的文件夹重命名
            if check_results['operator'] == 'submit':
                new_work_path = os.path.join(
                    var.storage, 
                    'temp', 
                    '{}_{}_{}'.format(
                        int(datetime.datetime.now().timestamp()), 
                        check_results['target'][1], 
                        '+'.join(json.loads(check_results['target'][10]).keys())
                    )
                )
            ## 对于完成审核，将文件夹移动至archive并重命名，重名时清除上一条记录
            if check_results['operator'] == 'finish':
                new_work_path = os.path.join(
                    var.storage, 
                    'archive', 
                    '{}'.format('+'.join(json.loads(check_results['target'][10]).keys()))
                )
                if os.path.isdir(new_work_path):
                    shutil.rmtree(new_work_path)
            ## 存档操作结束后更新check_results中的work_path
            shutil.move(check_results['work_path'], new_work_path)
            check_results['work_path'] = new_work_path
            logger.info('new work_path: {}'.format(new_work_path))

            # 发送邮件
            check_results['notification'] = {}
            ## 压缩附件目录
            archive_path = archive.archive(
                os.path.join(check_results['work_path'], 'attachments'), 
                '+'.join(json.loads(check_results['target'][10]).keys())
            )

            if check_results['operator'] == 'submit':
                check_results['notification']['mail'] = notification.build_submit_mail(
                    check_results['target'],
                    check_results['warnings']
                )
                mail.send(
                    check_results['target'][3],
                    check_results['notification']['mail']['subject'],
                    check_results['notification']['mail']['content'],
                    archive_path,
                    to_stdout=var.debug
                )
            if check_results['operator'] == 'finish':
                check_results['notification']['mail'] = notification.build_finish_mail(
                    check_results['target'],
                    check_results['warnings']
                )
                mail.send(
                    check_results['target'][1],
                    check_results['notification']['mail']['subject'],
                    check_results['notification']['mail']['content'],
                    archive_path,
                    needs_cc=True,
                    to_stdout=var.debug
                )
            os.remove(archive_path)        

            # 发送钉钉通知
            ## 钉钉机器人在初始化时无法校验参数是否正确，如果参数有误，dingtalkchatbot会抛出异常
            ## 钉钉通知是可选功能，报错时应当输出信息并继续
            try:
                if check_results['operator'] == 'submit':
                    check_results['notification']['dingtalk'] = notification.build_submit_dingtalk(
                        check_results['target'],
                        check_results['warnings']
                    )
                    dingtalk.send_markdown(
                        check_results['notification']['dingtalk']['subject'], 
                        check_results['notification']['dingtalk']['content'], 
                        mysql.t_user.fetch(check_results['target'][3])[2],
                        to_stdout=var.debug
                    )
                if check_results['operator'] == 'finish':
                    check_results['notification']['dingtalk'] = notification.build_finish_dingtalk(
                        check_results['target'],
                        check_results['warnings']
                    )
                    dingtalk.send_markdown(
                        check_results['notification']['dingtalk']['subject'], 
                        check_results['notification']['dingtalk']['content'], 
                        mysql.t_user.fetch(check_results['target'][1])[2],
                        to_stdout=var.debug
                    )
            except:
                # 向RM.log写入发送失败日志
                logger.error('dingtalk notification error', exc_info=True)

            # 发送企业微信通知
            try:
                if check_results['operator'] == 'submit':
                    check_results['notification']['wxwork'] = notification.build_submit_wxwork(
                        check_results['target'],
                        check_results['warnings']
                    )
                    wxwork.send_text(
                        check_results['notification']['wxwork']['content'], 
                        [check_results['target'][1], check_results['target'][3]],
                        to_stdout=var.debug
                    )
                if check_results['operator'] == 'finish':
                    check_results['notification']['wxwork'] = notification.build_finish_wxwork(
                        check_results['target'],
                        check_results['warnings']
                    )
                    wxwork.send_text(
                        check_results['notification']['wxwork']['content'], 
                        [check_results['target'][1], check_results['target'][3]],
                        to_stdout=var.debug
                    )
            except:
                # 向RM.log写入发送失败日志
                logger.error('wxwork notification error', exc_info=True)

            # 流程结束，向log_mail写入操作成功日志
            mysql.t_log.add(check_results)

        # 目前使用了单个异常捕获，任何异常都应在数据库和文件中分别记录日志
        except Exception as err:
            # 向log_mail写入操作失败日志
            mysql.t_log.add(check_results, '{}'.format(err))
            # 向RM.log写入操作失败日志
            logger.critical('mail failed ({})'.format(check_results['content']['user_id']), exc_info=True)
            # 向钉钉调试群发送操作失败告警
            dingtalk.send_markdown(
                '[RM] 处理异常', '({}) {}'.format(check_results['content']['user_id'], err),
                to_debug=True,
                to_stdout=var.debug
            )
            wxwork.send_text(
                '({}) {}'.format(check_results['content']['user_id'], err),
                to_debug=True,
                to_stdout=var.debug
            )


def do_resend(target:tuple, redirect:str=''):
    ''' 根据{target}的信息及类型（提交审核/完成审核），重新向任务相关人员发送邮件。

    Args:
        target(tuple): 项目记录（同主流程的target）
        redirect(str): 忽略原发件对象，将邮件发送至指定user_id

    Raises:
        AssertionError: 如果参数类型非法
    '''
    logger = logging.getLogger(__name__)
    logger.debug('args: {}'.format({'target': target, 'redirect': redirect}))
    assert type(target) == tuple and len(target) == 11, 'invalid arg: target'
    assert type(redirect) == str and mysql.t_user.__contains__(redirect), 'invalid arg: redirect'

    # 搜索最新的文件记录
    codes = list(json.loads(target[10]).keys())
    codes.sort()
    work_path = list(dir_paths(filtered_walk(
        os.path.join(var.storage, 'temp' if type(target[0]) == str else 'archive'), 
        included_dirs=['*{}'.format('+'.join(codes))], 
        depth=1, 
        min_depth=1
    )))[-1]
    logger.debug('found path: {}'.format(work_path))

    # 对于current中获取的target，重发[分配审核]
    if type(target[0]) == str:
        resend_notification = notification.build_submit_mail(target, [])
        resend_notification['subject'] = '(resend) ' + resend_notification['subject']
        to = redirect if redirect else target[3]
        logger.info('resending "{}" (分配审核) to "{}"'.format('+'.join(codes), to))
    # 对于history中获取的target，重发[完成审核]
    if type(target[0]) == int:
        resend_notification = notification.build_finish_mail(target, [])
        resend_notification['subject'] = '(resend) ' + resend_notification['subject']
        to = redirect if redirect else target[1]
        logger.info('resending "{}" (完成审核) to "{}"'.format('+'.join(codes), to))

    # 压缩附件
    archive_path = archive.archive(work_path, '+'.join(json.loads(target[10]).keys()))
    # 发送并清理临时文件
    mail.send(
        to,
        resend_notification['subject'],
        resend_notification['content'],
        archive_path,
        to_stdout=var.debug
    )
    os.remove(archive_path)


def do_clean():
    ''' 清理temp中的超时缓存。
    '''
    logger = logging.getLogger(__name__)

    # 删除temp中超过7天的文件夹
    for path in dir_paths(filtered_walk(os.path.join(var.storage, 'temp'), depth=1, min_depth=1)):
        if int(os.path.basename(path).split('_', 1)[0]) < datetime.datetime.now().timestamp() - 604800:
            logger.info('removed "{}" due to "expiration"'.format(path))
            shutil.rmtree(path)


def main_queue(q:multiprocessing.Queue):
    ''' 消息队列建议运行在独立的进程中，确保同时只有一个主逻辑运行。

    入队列item定义为{'command': xxx, 'kwargs': {xxx}}
    '''
    logger = logging.getLogger(__name__)
    logger.info('queue inited (PID: {})'.format(os.getpid()))

    while True:
        # 默认阻塞当前进程，直到队列中出现可用的对象
        item = q.get()
        logger.debug('new item "{}"'.format(item['command']))
        try:
            if item['command'] == 'attend':
                do_attend()
                do_clean()
                mysql.t_user.reset_status()
            elif item['command'] == 'mail':
                do_mail(item['kwargs'])
            elif item['command'] == 'resend':
                do_resend(item['kwargs']['target'], item['kwargs']['to'])
            else:
                pass
        except:
            pass

# def do_notify_archive():
#     ''' 向项目组长发送提醒归档的邮件。
#     '''

#     ret = mysql.t_history.analysis_procedure_1()
#     for row in ret:
#         mail = notification.build_monthly_projects_mail(row['authorid'], row['names'])
#         var.smtp_client.send(row['authorid'], mail['subject'], mail['content'])
