# -*- coding: UTF-8 -*-
import os
import logging, logging.config
import shutil
import datetime
import json
from walkdir import filtered_walk, file_paths, dir_paths
from docx import Document

import RM


def init():
    from configparser import ConfigParser
    config = ConfigParser()
    config.read(os.path.join('conf','RM.conf'), encoding='UTF-8')

    # ---运行模式（debug）---
    global debug
    debug = config.getboolean('mode', 'debug', fallback=False)

    # ---外部存储---
    global storage
    storage = config.get('path', 'storage', fallback='storage')
    ## 自动创建目录结构
    for check_dir in [os.path.join(storage, child_dir) for child_dir in ['temp', 'archive']]:
        if os.path.isdir(check_dir):
            continue
        os.mkdir(check_dir)

    # ---logger---
    dict_config = {
        'version': 1,
        'formatters': {
            'main': {
                'format': '%(asctime)s - %(levelname)s - %(name)s/%(funcName)s:%(lineno)d -> %(message)s',
                'datefmt': '%Y-%m-%d %H:%M:%S',
                # 'style': '%',
                # 'validate': True,
            },
        },
        # 'filters': {},
        'handlers': {
            'console': {
                'class' : 'logging.StreamHandler',
                'formatter': 'main',
                'level': 'DEBUG',
                # 'filters': '',
                'stream': 'ext://sys.stdout',
            },
        },
        'loggers': { 
            '': {
                'level': 'DEBUG',
                'propagate': False,
                # 'filters': [],
                'handlers': ['console'],
            },
        },
    }
    if config.get('log', 'dest', fallback=''):
        dict_config['handlers']['file'] = {
            'class' : 'logging.handlers.RotatingFileHandler',
            'formatter': 'main',
            'level': config.get('log', 'level', fallback='INFO').upper(),
            # 'filters': '',
            'filename': os.path.join(storage, os.path.basename(config.get('log', 'dest', fallback=''))),
            'maxBytes': 204800,
            'backupCount': 5,
        }
        dict_config['loggers']['']['handlers'].append('file')
    logging.config.dictConfig(dict_config)

    logger = logging.getLogger(__name__)
    logger.info('logging to "{}" ({})'.format(
        dict_config['handlers']['file']['filename'], 
        dict_config['handlers']['file']['level'],
    ))

    # ---mysql---
    RM.mysql.init(
        user=config.get('mysql', 'user', fallback='rm'), 
        password=config.get('mysql', 'pass', fallback='rm'), 
        host=config.get('mysql', 'host', fallback='127.0.0.1'), 
        database=config.get('mysql', 'db', fallback='rm'),
        port=config.getint('mysql', 'port', fallback=3306),
    )

    # ---redis---
    global stream
    stream = RM.RedisStream(
        host=config.get('redis', 'host', fallback='127.0.0.1'), 
        password=config.get('redis', 'pass', fallback='rm'), 
    )

    # ---mail---
    global mail
    pop3_config = {
        'username': config.get('pop3', 'user', fallback='rm@example.com'), 
        'password': config.get('pop3', 'pass', fallback='rm'), 
        'host': config.get('pop3', 'host', fallback='example.com'), 
        'port': config.getint('pop3', 'port', fallback=110),
        'ssl': config.getboolean('pop3', 'ssl', fallback=False),
        'tls': config.getboolean('pop3', 'tls', fallback=False),
    }
    smtp_config = {
        'username': config.get('smtp', 'user', fallback='rm@example.com'), 
        'password': config.get('smtp', 'pass', fallback='rm'), 
        'host': config.get('smtp', 'host', fallback='example.com'), 
        'port': config.getint('smtp', 'port', fallback=25),
        'ssl': config.getboolean('smtp', 'ssl', fallback=False),
        'tls': config.getboolean('smtp', 'tls', fallback=False),
    }
    mail_config = {
        'default_domain': config.get('mail', 'domain', fallback='example.com'),
        'default_cc': config.get('mail', 'manager', fallback=''),
    }
    mail = RM.Mail(pop3_config, smtp_config, mail_config)

    # ---archive---
    global archive
    bin_path = {
        'winrar': config.get('archive', 'winrar_bin', fallback='C:\\Program Files\\WinRAR\\WinRAR.exe'),
        'rar': config.get('archive', 'rar_bin', fallback='rar'),
        'unrar': config.get('archive', 'unrar_bin', fallback='unrar'),
        'unar': config.get('archive', 'unar_bin', fallback='unar')
    }
    archive = RM.Archive(bin_path, config.get('archive', 'pass', fallback=''))

    # ---dingtalk---
    global dingtalk
    chatbot = {
        'webhook': config.get('dingtalk', 'webhook', fallback=''),
        'secret': config.get('dingtalk', 'secret', fallback=''),
    }
    chatbot_debug = {
        'webhook': config.get('dingtalk', 'webhook_debug', fallback=''),
        'secret': config.get('dingtalk', 'secret_debug', fallback='')
    }
    dingtalk = RM.Dingtalk(
        chatbot, 
        chatbot_debug,
        config.get('dingtalk', 'attend', fallback=''),
        config.get('dingtalk', 'interaction', fallback=''),
        config.getboolean('dingtalk', 'enable', fallback=False)
    )

    # ---wxwork---
    global wxwork
    wxwork = RM.WXWork(
        config.get('wxwork', 'corpid', fallback=''),
        config.get('wxwork', 'agentid', fallback=''),
        config.get('wxwork', 'secret', fallback=''),
        config.get('wxwork', 'admin_userid', fallback=''),
        config.getboolean('wxwork', 'enable', fallback=False)
    )


def do_mail(check_results:dict=None):
    ''' 邮件管理函数的主逻辑，对于传入的单个邮件，依次进行邮件检查、数据库操作、发送通知步骤。

    传入的{keywords}用于自定义处理逻辑。

    Args:
        keywords(dict): 拉取邮件时使用的关键词

    Returns:
        处理邮件数量
    '''
    logger = logging.getLogger(__name__)
    logger.debug('args: {}'.format({'check_results': check_results}))

    logger.info('reading "{}" from "{}"'.format(check_results['operator'], check_results['mail']['from']))
    try:
        # Step 1: 检查邮件内容及附件内容，得到work_path和所有相关信息，生成需要的文件并清理文件夹
        attachments_path = os.path.join(check_results['work_path'], 'attachments')

        # content的常规报错情况为：
        #   发件人非法
        ret = RM.validator.check_mail_content(check_results['mail'])
        check_results['content'] = ret['content']
        check_results['warnings'] += ret['warnings']

        # 解压工作目录中的所有压缩包，并添加解压失败的告警
        for archive_file_path_with_error in archive.extract(attachments_path):
            ret['warnings'].append('解压失败："{}"'.format(os.path.basename(archive_file_path_with_error)))
        # attachment的常规报错情况为：
        #   未从附件中读取到有效文档
        ret = RM.validator.check_mail_attachment(attachments_path, check_results['operator'])
        check_results['attachment'] = ret['attachment']
        check_results['warnings'] += ret['warnings']

        # 提交审核时，生成XT13
        if check_results['operator'] == 'submit':
            for code in check_results['attachment']['codes']:
                try:
                    logger.info('generating XT13 for "{}"'.format(code))
                    # 根据项目编号读取不同的审核意见单模板
                    if 'PRO' in code or 'PST' in code:
                        template_document = Document(os.path.join('template', 'RD-XT13测评报告审核、签发意见单-PROPST.docx'))
                    elif 'SOF' in code:
                        template_document = Document(os.path.join('template', 'RD-XT13测评报告审核、签发意见单-SOF.docx'))
                    elif 'DSYS' in code:
                        template_document = Document(os.path.join('template', 'RD-XT13测评报告审核、签发意见单-DSYS.docx'))
                    elif 'SRV' in code:
                        template_document = Document(os.path.join('template', 'RD-XT13测评报告审核、签发意见单-SRV.docx'))
                    elif 'PER' in code:
                        template_document = Document(os.path.join('template', 'RD-XT13测评报告审核、签发意见单-PER.docx'))
                    elif 'FUN' in code:
                        template_document = Document(os.path.join('template', 'RD-XT13测评报告审核、签发意见单-FUN.docx'))
                    else:
                        # 目前无视其他类型的报告
                        continue 

                    # 项目编号(bold)
                    template_document.paragraphs[0].add_run(code).bold = True
                    # 项目名称
                    template_document.tables[0].cell(1, 1).text = check_results['attachment']['names'].get(code, '')
                    # 报告撰写人
                    template_document.tables[0].cell(2, 1).text = check_results['content']['name']

                    # 另存为
                    template_document.save(os.path.join(attachments_path, 'RD-XT13测评报告审核、签发意见单{}.docx'.format(code)))
                except Exception as err:
                    logger.warning('generating XT13 failed: {}'.format(err))
                    check_results['warnings'].append('generating XT13 failed for "{}"'.format(code))

        # 完成审核时，删除多余文件
        if check_results['operator'] == 'finish':
            for file_path in file_paths(filtered_walk(attachments_path, excluded_files=['*.doc', '*.docx', '*.rar', '*.zip', '*.7z', '*.tar'])):
                os.remove(file_path)
                logger.info('removed "{}"'.format(os.path.basename(file_path)))


        # Step 2: 进行数据库操作，得到操作结果target

        # 提交审核的情况
        if check_results['operator'] == 'submit':
            # 向current库中插入该项目，add接口只需要传入项目基本信息
            current_id = RM.mysql.t_current.add(
                check_results['attachment']['names'], 
                check_results['attachment']['company'],
                check_results['attachment']['pages'],
                check_results['content']['urgent'],
                check_results['content']['user_id'], 
                check_results['mail']['timestamp'] 
            )

            # 使用edit接口分配项目
            ## 如果有指定审核人，直接设置为该人（前一步骤已经校验了有效性）
            if check_results['content']['force']:
                RM.mysql.t_current.edit(current_id, reviewerid=check_results['content']['force'])
            ## 否则调用RM.mysql.t_user.pop获取下一顺位的审核人，传入的excludes拼接了组员和自己
            else:
                RM.mysql.t_current.edit(
                    current_id, 
                    reviewerid=RM.mysql.t_user.pop(
                        excludes=check_results['content']['excludes'] + [check_results['content']['user_id']],
                        urgent=check_results['content']['urgent'],
                    )[0][0]
                )

            # 从current中取回操作结果，写入target
            check_results['target'] = RM.mysql.t_current.fetch(current_id)
            logger.info('(submit) "{}" -> "{}"'.format(check_results['target'][1], check_results['target'][3]))

        # 完成审核的情况
        if check_results['operator'] == 'finish':
            codes = '+'.join(check_results['attachment']['codes'])
            ret = RM.mysql.t_current.search('+'.join(check_results['attachment']['codes']))['all']
            assert len(ret) == 1, 'Not unique result for code: "{}".'.format(codes)
            # 从current库中正常删除项目，delete接口会自动将结果写入history
            RM.mysql.t_current.delete(ret[0][0], check_results['mail']['timestamp'])

            # 从history中取回操作结果，写入target
            ## 由于插入history时不检测唯一性，当用户输错项目编号时，search结果可能重复。
            ## 在此取第一个结果作为有效返回
            check_results['target'] = RM.mysql.t_history.search(codes)['all'][0]
            logger.info('(finish) "{}" <- "{}"'.format(check_results['target'][1], check_results['target'][3]))
        
        logger.debug('target: {}'.format(check_results['target']))


        # Step 3: 存档work_path，并发送通知消息
        # 该步骤主要用于生成通知内容并发送通知。邮件发送失败将抛出异常，钉钉通知发送失败将告警并继续处理

        ## 去除所有子目录结构
        for file_path in file_paths(filtered_walk(attachments_path, min_depth=1)):
            shutil.copy(file_path, attachments_path)
        for dir_path in dir_paths(filtered_walk(attachments_path, min_depth=1, depth=1)):
            shutil.rmtree(dir_path)

        ## 对于提交审核，将attachments_path重命名并移动至temp中
        if check_results['operator'] == 'submit':
            new_work_path = os.path.join(
                storage, 
                'temp', 
                '{}_{}_{}'.format(
                    datetime.datetime.now().timestamp(), 
                    check_results['target'][1], 
                    '+'.join(json.loads(check_results['target'][10]))
                )
            )
        ## 对于完成审核，将attachments_path重命名并移动至archive中，重名时清除上一条记录，同时删除temp中的提交审核记录
        if check_results['operator'] == 'finish':
            new_work_path = os.path.join(
                storage, 
                'archive', 
                '+'.join(json.loads(check_results['target'][10])),
            )
            if os.path.isdir(new_work_path):
                shutil.rmtree(new_work_path)
            for dir_path in dir_paths(filtered_walk(
                os.path.join(storage, 'temp'), 
                included_dirs=['*'+'+'.join(json.loads(check_results['target'][10]))], 
                depth=1, 
                min_depth=1,
            )):
                shutil.rmtree(dir_path)
        ## 存档操作结束后更新check_results中的work_path
        shutil.move(attachments_path, new_work_path)
        shutil.rmtree(check_results['work_path'])
        check_results['work_path'] = new_work_path
        logger.info('new work_path: {}'.format(new_work_path))

        # 发送邮件
        check_results['notification'] = {}
        ## 压缩附件目录
        archive_path = archive.archive(
            check_results['work_path'], 
            '+'.join(json.loads(check_results['target'][10]))
        )

        if check_results['operator'] == 'submit':
            check_results['notification']['mail'] = RM.notification.build_submit_mail(
                check_results['target'],
                check_results['warnings']
            )
            mail.send(
                check_results['target'][3],
                check_results['notification']['mail']['subject'],
                check_results['notification']['mail']['content'],
                archive_path,
                to_stdout=debug
            )
        if check_results['operator'] == 'finish':
            check_results['notification']['mail'] = RM.notification.build_finish_mail(
                check_results['target'],
                check_results['warnings']
            )
            mail.send(
                check_results['target'][1],
                check_results['notification']['mail']['subject'],
                check_results['notification']['mail']['content'],
                archive_path,
                needs_cc=True,
                to_stdout=debug
            )
        os.remove(archive_path)        

        # 发送钉钉通知
        ## 钉钉机器人在初始化时无法校验参数是否正确，如果参数有误，dingtalkchatbot会抛出异常
        ## 钉钉通知是可选功能，报错时应当输出信息并继续
        try:
            if check_results['operator'] == 'submit':
                check_results['notification']['dingtalk'] = RM.notification.build_submit_dingtalk(
                    check_results['target'],
                    check_results['warnings']
                )
                dingtalk.send_markdown(
                    check_results['notification']['dingtalk']['subject'], 
                    check_results['notification']['dingtalk']['content'], 
                    RM.mysql.t_user.fetch(check_results['target'][3])[2],
                    to_stdout=debug
                )
            if check_results['operator'] == 'finish':
                check_results['notification']['dingtalk'] = RM.notification.build_finish_dingtalk(
                    check_results['target'],
                    check_results['warnings']
                )
                dingtalk.send_markdown(
                    check_results['notification']['dingtalk']['subject'], 
                    check_results['notification']['dingtalk']['content'], 
                    RM.mysql.t_user.fetch(check_results['target'][1])[2],
                    to_stdout=debug
                )
        except:
            # 向RM.log写入发送失败日志
            logger.error('dingtalk notification error', exc_info=True)

        # 发送企业微信通知
        try:
            if check_results['operator'] == 'submit':
                check_results['notification']['wxwork'] = RM.notification.build_submit_wxwork(
                    check_results['target'],
                    check_results['warnings']
                )
                wxwork.send_text(
                    check_results['notification']['wxwork']['content'], 
                    [check_results['target'][1], check_results['target'][3]],
                    to_stdout=debug
                )
            if check_results['operator'] == 'finish':
                check_results['notification']['wxwork'] = RM.notification.build_finish_wxwork(
                    check_results['target'],
                    check_results['warnings']
                )
                wxwork.send_text(
                    check_results['notification']['wxwork']['content'], 
                    [check_results['target'][1], check_results['target'][3]],
                    to_stdout=debug
                )
        except:
            # 向RM.log写入发送失败日志
            logger.error('wxwork notification error', exc_info=True)

        # 流程结束，向log_mail写入操作成功日志
        RM.mysql.t_log.add(check_results)

    # 目前使用了单个异常捕获，任何异常都应在数据库和文件中分别记录日志
    except Exception as err:
        # 向log_mail写入操作失败日志
        RM.mysql.t_log.add(check_results, '{}'.format(err))
        # 向RM.log写入操作失败日志
        logger.critical('mail failed ({})'.format(check_results['content']['user_id']), exc_info=True)
        # 向钉钉调试群发送操作失败告警
        dingtalk.send_markdown(
            '[RM] 处理异常', '({}) {}'.format(check_results['content']['user_id'], err),
            to_debug=True,
            to_stdout=debug
        )
        wxwork.send_text(
            '({}) {}'.format(check_results['content']['user_id'], err),
            to_debug=True,
            to_stdout=debug
        )


def do_resend(id:str|int, redirect:str=''):
    ''' 根据{id}的信息及类型（提交审核/完成审核），重新向任务相关人员发送邮件。

    Args:
        id(tuple): 项目记录ID
        redirect(str): 忽略原发件对象，将邮件发送至指定user_id

    Raises:
        AssertionError: 如果参数类型非法
    '''
    logger = logging.getLogger(__name__)
    logger.debug('args: {}'.format({'id': id, 'redirect': redirect}))
    if type(id) == int:
        target = RM.mysql.t_history.fetch(id)
    else:
        target = RM.mysql.t_current.fetch(id)
    if not type(redirect) == str or not RM.mysql.t_user.__contains__(redirect):
        logger.warning('invalid arg: redirect')
        redirect = ''

    # 搜索最新的文件记录
    codes = list(json.loads(target[10]))
    codes.sort()
    work_path = list(dir_paths(filtered_walk(
        os.path.join(storage, 'temp' if type(target[0]) == str else 'archive'), 
        included_dirs=['*{}'.format('+'.join(codes))], 
        depth=1, 
        min_depth=1
    )))[-1]
    logger.debug('found path: {}'.format(work_path))

    # 对于current中获取的target，重发[分配审核]
    if type(target[0]) == str:
        resend_notification = RM.notification.build_submit_mail(target, [])
        resend_notification['subject'] = '(resend) ' + resend_notification['subject']
        to = redirect if redirect else target[3]
        logger.info('resending "{}" (分配审核) to "{}"'.format('+'.join(codes), to))
    # 对于history中获取的target，重发[完成审核]
    if type(target[0]) == int:
        resend_notification = RM.notification.build_finish_mail(target, [])
        resend_notification['subject'] = '(resend) ' + resend_notification['subject']
        to = redirect if redirect else target[1]
        logger.info('resending "{}" (完成审核) to "{}"'.format('+'.join(codes), to))

    # 压缩附件
    archive_path = archive.archive(work_path, '+'.join(json.loads(target[10])))
    # 发送并清理临时文件
    mail.send(
        to,
        resend_notification['subject'],
        resend_notification['content'],
        archive_path,
        to_stdout=debug
    )
    os.remove(archive_path)


if __name__ == "__main__":
    import socket
    init()
    logger = logging.getLogger('main')
    logger.warning('worker "{}" initiated'.format(socket.gethostname())) 

    while True:
        # 默认阻塞当前进程，直到队列中出现可用的对象
        logger.debug('waiting stream')
        item = stream.read()
        logger.debug('new item "{}"'.format(item['command']))
        try:
            if item['command'] == 'receive':
                received_mails = mail.receive(os.path.join(storage, 'temp'), item['kwargs'])
                for received_mail in received_mails:
                    do_mail(received_mail)
            elif item['command'] == 'read':
                local_mail = mail.read(
                    os.path.join(storage, 'temp'),
                    item['kwargs']['key'],
                    item['kwargs']['file'],
                )
                do_mail(local_mail)
            elif item['command'] == 'resend':
                do_resend(**item['kwargs'])
        except Exception as err:
            logger.error(err, exc_info=True)